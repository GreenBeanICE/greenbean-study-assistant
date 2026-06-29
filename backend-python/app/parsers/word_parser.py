"""Word 文档解析器，按标题/段落/表格产出统一节点。"""
import io
from typing import List, Dict, Any
from docx import Document
from app.utils.text_utils import clean_text


class WordParser:
    """解析 .docx 格式的 Word 文档，提取结构化文本内容。"""

    @staticmethod
    def _heading_metadata(para) -> list[dict[str, Any]]:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            return [{
                "level": int(para.style.name.replace("Heading ", "0") or "0"),
                "text": para.text.strip(),
            }]
        return []

    @staticmethod
    def _make_node(content: str, sequence_index: int, *, headings: list[dict[str, Any]] | None = None,
                   block_type: str = "paragraph", paragraphs_count: int = 0, tables_count: int = 0) -> Dict[str, Any]:
        normalized_content = clean_text(content)
        return {
            "page_number": None,
            "content": normalized_content,
            "char_count": len(normalized_content),
            "parser_name": "WordParser",
            "parser_version": "1.0.0",
            "metadata": {
                "source_type": "word",
                "paragraphs_count": paragraphs_count,
                "tables_count": tables_count,
                "headings": headings or [],
                "block_type": block_type,
                "sequence_index": sequence_index,
            },
        }

    def parse(self, file_content: bytes) -> List[Dict[str, Any]]:
        """解析 Word 文档字节流，按块构建统一节点。"""
        doc = Document(io.BytesIO(file_content))

        nodes: list[dict[str, Any]] = []
        sequence_index = 0
        non_empty_paragraphs = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            non_empty_paragraphs += 1
            nodes.append(
                self._make_node(
                    text,
                    sequence_index,
                    headings=self._heading_metadata(para),
                    block_type="heading" if self._heading_metadata(para) else "paragraph",
                    paragraphs_count=1,
                )
            )
            sequence_index += 1

        table_count = 0
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_content.append(" | ".join(row_cells))
            table_count += 1
            nodes.append(
                self._make_node(
                    "\n".join(table_content),
                    sequence_index,
                    block_type="table",
                    tables_count=1,
                )
            )
            sequence_index += 1

        if nodes:
            return nodes

        return [
            self._make_node(
                "",
                0,
                paragraphs_count=non_empty_paragraphs,
                tables_count=table_count,
            )
        ]
